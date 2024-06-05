import os
import pdfplumber
import streamlit as st
from docx import Document
from docx.shared import Pt
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings, HarmBlockThreshold, HarmCategory
from langchain.chains.question_answering import load_qa_chain
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv


load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key = GOOGLE_API_KEY)

chain_gemini_pro = ChatGoogleGenerativeAI(model="gemini-pro", safety_settings=
            {HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE}, 
            temperature=0, google_api_key=GOOGLE_API_KEY)


def extract_text_from_pdf(pdf_docs):
    text = ""
    try:
        for pdf in pdf_docs:
            with pdfplumber.open(pdf) as pdf_file:
                for page in pdf_file.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
    except Exception as e:
        st.error(f"Failed to read PDF files: {e}")
    return text

def get_text_chunks(text):
    try:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=100)
        chunks = text_splitter.split_text(text)
    except Exception as e:
        st.error(f"Error splitting text into chunks: {e}")
        return []
    return chunks

def create_vector_store(text_chunks):
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        if os.path.exists("faiss_index"):
            vector_store = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            new_index = FAISS.from_texts(text_chunks, embedding=embeddings)
            vector_store.merge_from(new_index)
            vector_store.save_local("faiss_index")
        else:
            vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
            vector_store.save_local("faiss_index")
    except Exception as e:
        st.error(f"Failed to create or update vector store: {e}")

def initialize_conversational_chain():
    try:
        prompt_template = """
        Answer the question as detailed as possible from the provided context, and always start with some side heading if needed(make sure to get it from question or context), 
        make sure to provide all the details, if the answer is contains mutiple steps generate point wise output,   
        if the answer is not in provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
        Context:\n {context}?\n
        Question: \n{question}\n
        Answer:
        """
        model = ChatGoogleGenerativeAI(model="gemini-pro", safety_settings={
                HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
            },temperature=0, google_api_key=GOOGLE_API_KEY)

        prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
        chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    except Exception as e:
        st.error(f"Failed to initialize conversational chain: {e}")
        return None
    return chain

def get_user_input(user_question):
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        docs = new_db.similarity_search(user_question)
        chain = initialize_conversational_chain()
        response = chain({"input_documents":docs, "question": user_question}, return_only_outputs=True)
    except Exception as e:
        st.error(f"Error fetching user input: {e}")
        return "Error: Unable to process the request."
    return response["output_text"]

def download_chat_history():
    try:
        document = Document()
        document.add_heading('Chat History', level=1)
        for user_input, bot_response in zip(st.session_state["past"], st.session_state["generated"]):
            user_paragraph = document.add_paragraph()
            user_paragraph.add_run("You: ").bold = True
            user_paragraph.add_run(user_input).font.size = Pt(12)
            bot_paragraph = document.add_paragraph()
            bot_paragraph.add_run("Bot: ").bold = True
            bot_paragraph.add_run(bot_response).font.size = Pt(12)
            document.add_paragraph()
        chat_history_file_path = "chat_history.docx"
        document.save(chat_history_file_path)
    except Exception as e:
        st.error(f"Failed to download chat history: {e}")
        return None
    return chat_history_file_path


def main():
    st.set_page_config("Chat PDFs")
    st.markdown("<h1 style='text-align: center;'>Talk to MRI</h1>", unsafe_allow_html=True)

    # Initialize session states for chat history if they do not already exist
    if 'past' not in st.session_state:
        st.session_state['past'] = []
    if 'generated' not in st.session_state:
        st.session_state['generated'] = []
    if 'current_input' not in st.session_state:
        st.session_state['current_input'] = ""  # Initialize a state to store current input

    # Create containers for chat history and user input
    response_container = st.container()
    container = st.container()

    with container:
        with st.form(key='my_form', clear_on_submit=True):
            # Use session state to manage the current input
            user_question = st.text_input("Query:", value=st.session_state['current_input'], placeholder="Ask a Question from the PDF Files", key='input')
            submit_button = st.form_submit_button(label="Send")

    if submit_button and user_question:
        output = get_user_input(user_question)  # Call the function to get a response
        st.session_state['past'].append(user_question)
        st.session_state['generated'].append(output)
        st.session_state['current_input'] = ""  # Reset the current input after processing

    # Render the chat history
    with response_container:
        for query, response in zip(st.session_state['past'], st.session_state['generated']):
            with st.chat_message("user"):
                st.write(query)  # Changed from user_input to query to avoid name conflict
            with st.chat_message("assistant"):
                st.write(response)  # Changed from bot_response to response for consistency

    # Download button placed outside of the chat update logic
    with container:
        if st.button("Download Chat History"):
            chat_history_path = download_chat_history()
            with open(chat_history_path, "rb") as file:
                st.download_button(
                    label="Download Chat History",
                    data=file,
                    file_name="chat_history.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    with st.sidebar:
        st.title("Menu:")
        pdf_docs = st.file_uploader("Upload your PDF Files and Click on the Submit & Process Button", accept_multiple_files=True)
        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                raw_text = extract_text_from_pdf(pdf_docs)
                text_chunks = get_text_chunks(raw_text)
                create_vector_store(text_chunks)
                st.success("Done")

if __name__ == "__main__":
    main()
